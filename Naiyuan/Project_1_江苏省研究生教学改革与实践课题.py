import os
import pandas as pd
import docx
from docx import Document
import xlrd
from xlutils.copy import copy
import openpyxl
from openpyxl import load_workbook

Result=pd.read_excel('江苏省研究生教育教学改革研究与实践课题.xlsx')
data1=pd.read_excel('2016年江苏省研究生教育教学改革研究与实践课题名单.xls')
data2=Document('2017年江苏省研究生教育教学改革课题.docx')
data3=Document('2018年江苏省研究生教育教学改革课题.docx')
data4=Document('2019年江苏省研究生教育教学改革课题名单.docx')

#################################2016.xls
number_of_project_2016=[]
name_of_institution_2016=[]
hoster_2016=[]
name_of_project_2016=[]
class_of_project_2016=[]
type_of_money_2016=[]

wb2=xlrd.open_workbook('2016年江苏省研究生教育教学改革研究与实践课题名单.xls')
sheet=wb2.sheet_by_name('教改课题')

for a in range(sheet.nrows):
    cells=sheet.row_values(a)
    data=str(cells[0])
    number_of_project_2016.append(data)
del number_of_project_2016[0]
del number_of_project_2016[0]
del number_of_project_2016[0]
del number_of_project_2016[0]

for a in range(sheet.nrows):
    cells=sheet.row_values(a)
    data=str(cells[1])
    name_of_project_2016.append(data)
del name_of_project_2016[0]
del name_of_project_2016[0]
del name_of_project_2016[0]
del name_of_project_2016[0]

for a in range(sheet.nrows):
    cells=sheet.row_values(a)
    data=str(cells[2])
    name_of_institution_2016.append(data)
del name_of_institution_2016[0]
del name_of_institution_2016[0]
del name_of_institution_2016[0]
del name_of_institution_2016[0]

for a in range(sheet.nrows):
    cells=sheet.row_values(a)
    data=str(cells[3])
    class_of_project_2016.append(data)
del class_of_project_2016[0]
del class_of_project_2016[0]
del class_of_project_2016[0]
del class_of_project_2016[0]

for a in range(sheet.nrows):
    cells=sheet.row_values(a)
    data=str(cells[4])
    type_of_money_2016.append(data)
del type_of_money_2016[0]
del type_of_money_2016[0]
del type_of_money_2016[0]
del type_of_money_2016[0]

for a in range(sheet.nrows):
    cells=sheet.row_values(a)
    data=str(cells[5])
    hoster_2016.append(data)
del hoster_2016[0]
del hoster_2016[0]
del hoster_2016[0]
del hoster_2016[0]


##############################2017年docx
number_of_project_2017=[]
name_of_institution_2017=[]
hoster_2017=[]
name_of_project_2017=[]
class_of_project_2017=[]
type_of_money_2017=[]

table2=data2.tables
for i in range(len(table2)):
    table=table2[i]
    #获取表格的行
    table_row=table.rows
    #获取每一行的值
    for i in range(len(table_row)):
        row_data=[]
        row_cells=table_row[i].cells
        #读取每一行单元格的内容
        for cell in row_cells:
            #单元格内容
            row_data.append(cell.text)
        # print(row_data)
        number_of_project_2017.append(row_data[0])
        name_of_institution_2017.append(row_data[1])
        hoster_2017.append(row_data[2])
        name_of_project_2017.append(row_data[3])
        class_of_project_2017.append(row_data[4])
        type_of_money_2017.append(row_data[5])

##############################2018年docx
number_of_project_2018=[]
name_of_institution_2018=[]
hoster_2018=[]
name_of_project_2018=[]
class_of_project_2018=[]
type_of_money_2018=[]

table3=data3.tables
for i in range(len(table3)):
    table=table3[i]
    #获取表格的行
    table_row=table.rows
    #获取每一行的值
    for i in range(len(table_row)):
        row_data=[]
        row_cells=table_row[i].cells
        #读取每一行单元格的内容
        for cell in row_cells:
            #单元格内容
            row_data.append(cell.text)
        # print(row_data)
        number_of_project_2018.append(row_data[0])
        name_of_institution_2018.append(row_data[1])
        hoster_2018.append(row_data[2])
        name_of_project_2018.append(row_data[3])
        class_of_project_2018.append(row_data[4])
        type_of_money_2018.append(row_data[5])

##############################2019年docx
number_of_project_2019=[]
name_of_institution_2019=[]
hoster_2019=[]
name_of_project_2019=[]
class_of_project_2019=[]
type_of_money_2019=[]

table4=data4.tables
for i in range(len(table4)):
    table=table4[i]
    #获取表格的行
    table_row=table.rows
    #获取每一行的值
    for i in range(len(table_row)):
        row_data=[]
        row_cells=table_row[i].cells
        #读取每一行单元格的内容
        for cell in row_cells:
            #单元格内容
            row_data.append(cell.text)
        # print(row_data)
        number_of_project_2019.append(row_data[0])
        name_of_institution_2019.append(row_data[1])
        hoster_2019.append(row_data[2])
        name_of_project_2019.append(row_data[3])
        class_of_project_2019.append(row_data[4])
        type_of_money_2019.append(row_data[5])


###################写文件
print(len(number_of_project_2016))
print(len(number_of_project_2017))
print(len(number_of_project_2018))
print(len(number_of_project_2019))

wb2=load_workbook("江苏省研究生教育教学改革研究与实践课题.xlsx")
wb3=wb2.active

for i in range(len(number_of_project_2016)):
    wb3.cell(i+2,8,'2016')
    wb3.cell(i+2,1,number_of_project_2016[i])
    wb3.cell(i+2,2,name_of_institution_2016[i])
    wb3.cell(i+2,3,hoster_2016[i])
    wb3.cell(i+2,4,name_of_project_2016[i])
    wb3.cell(i+2,5,number_of_project_2016[i])
    wb3.cell(i+2,6,class_of_project_2016[i])
    wb3.cell(i+2,7,type_of_money_2016[i])
for i in range(len(number_of_project_2017)):
    wb3.cell(i+227,8,'2017')
    wb3.cell(i + 227, 1, number_of_project_2017[i])
    wb3.cell(i + 227, 2, name_of_institution_2017[i])
    wb3.cell(i + 227, 3, hoster_2017[i])
    wb3.cell(i + 227, 4, name_of_project_2017[i])
    wb3.cell(i + 227, 5, number_of_project_2017[i])
    wb3.cell(i + 227, 6, class_of_project_2017[i])
    wb3.cell(i + 227, 7, type_of_money_2017[i])
for i in range(len(number_of_project_2018)):
    wb3.cell(i+440,8,'2018')
    wb3.cell(i + 440, 1, number_of_project_2018[i])
    wb3.cell(i + 440, 2, name_of_institution_2018[i])
    wb3.cell(i + 440, 3, hoster_2018[i])
    wb3.cell(i + 440, 4, name_of_project_2018[i])
    wb3.cell(i + 440, 5, number_of_project_2018[i])
    wb3.cell(i + 440, 6, class_of_project_2018[i])
    wb3.cell(i + 440, 7, type_of_money_2018[i])
for i in range(len(number_of_project_2019)):
    wb3.cell(i+704,8,'2019')
    wb3.cell(i + 704, 1, number_of_project_2019[i])
    wb3.cell(i + 704, 2, name_of_institution_2019[i])
    wb3.cell(i + 704, 3, hoster_2019[i])
    wb3.cell(i + 704, 4, name_of_project_2019[i])
    wb3.cell(i + 704, 5, number_of_project_2019[i])
    wb3.cell(i + 704, 6, class_of_project_2019[i])
    wb3.cell(i + 704, 7, type_of_money_2019[i])
wb2.save("江苏省研究生教育教学改革研究与实践课题.xlsx")